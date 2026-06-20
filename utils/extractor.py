import os
import io
import pdfplumber
import PyPDF2
from docx import Document
import logging
from typing import Union, BinaryIO

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ResumeExtractor:
    """
    Utility class to extract raw text from PDF and DOCX files.
    Supports both file paths (strings) and file-like objects (BytesIO/BinaryIO).
    """
    
    @staticmethod
    def extract_text(file_source: Union[str, BinaryIO, io.BytesIO], filename: str) -> str:
        """
        Extract text from a file-like object or string path based on extension.
        
        Args:
            file_source: File path (str) or file-like stream (BytesIO).
            filename: Name of the file, used to determine the file type extension.
            
        Returns:
            Extracted text content as a clean string.
        """
        ext = os.path.splitext(filename)[1].lower()
        
        # If file_source is a path string, open it
        if isinstance(file_source, str):
            if not os.path.exists(file_source):
                raise FileNotFoundError(f"File not found at: {file_source}")
            with open(file_source, 'rb') as f:
                file_bytes = io.BytesIO(f.read())
        else:
            # It's already a stream or bytes. Ensure it's in a seekable stream.
            if hasattr(file_source, 'read'):
                content = file_source.read()
                file_bytes = io.BytesIO(content)
                # Reset stream pointer of original file if possible
                if hasattr(file_source, 'seek'):
                    try:
                        file_source.seek(0)
                    except Exception:
                        pass
            else:
                raise ValueError("Invalid file source type. Must be a path string or file stream.")

        if ext == '.pdf':
            return ResumeExtractor.extract_from_pdf(file_bytes)
        elif ext in ['.docx', '.doc']:
            return ResumeExtractor.extract_from_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported file format: {ext}. Only PDF and DOCX are supported.")
            
    @staticmethod
    def extract_from_pdf(file_stream: io.BytesIO) -> str:
        """
        Extract text from PDF using pdfplumber, falling back to PyPDF2.
        """
        text = ""
        # 1. Try pdfplumber first (excellent layout parsing)
        try:
            file_stream.seek(0)
            with pdfplumber.open(file_stream) as pdf:
                pages_text = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)
                text = "\n".join(pages_text)
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}. Falling back to PyPDF2.")
            
        # 2. Fallback to PyPDF2 if pdfplumber returned empty or crashed
        if not text.strip():
            try:
                file_stream.seek(0)
                reader = PyPDF2.PdfReader(file_stream)
                pages_text = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)
                text = "\n".join(pages_text)
            except Exception as e:
                logger.error(f"PyPDF2 failed as well: {e}")
                
        return text.strip()

    @staticmethod
    def extract_from_docx(file_stream: io.BytesIO) -> str:
        """
        Extract text from a DOCX file stream using python-docx.
        """
        try:
            file_stream.seek(0)
            doc = Document(file_stream)
            full_text = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
                    
            # Extract tables content
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        full_text.append(" | ".join(row_cells))
                        
            return "\n".join(full_text).strip()
        except Exception as e:
            logger.error(f"python-docx failed: {e}")
            raise RuntimeError(f"Error parsing DOCX file: {e}")
